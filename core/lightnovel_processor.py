"""
Light Novel Processor for Game Translation Tool

Handles processing of light novels in various formats:
- TXT files (plain text)
- DOCX files (Microsoft Word documents)
- PDF files (Portable Document Format)
- EPUB files (Electronic Publication format with chapter support)
"""

import os
import re
import json
import chardet
import tiktoken
from typing import Dict, List, Tuple, Optional, Any
import logging

class LightNovelProcessor:
    def __init__(self):
        self.supported_extensions = ['.txt', '.docx', '.pdf', '.epub']
        self.output_formats = ['.txt', '.docx', '.pdf', '.epub']
        
        # Import optional dependencies
        self.docx_available = self._check_docx_support()
        self.pdf_available = self._check_pdf_support()
        self.epub_available = self._check_epub_support()
        
        # Eroge detection patterns
        self.eroge_keywords = [
            # Japanese erotic terms
            'エッチ', '性的', '裸', '胸', '乳房', 'おっぱい', 'セックス', '愛液', 
            '挿入', '射精', '絶頂', 'オーガズム', '感じる', '濡れ', 'キス',
            '抱く', '犯す', '責める', '舐める', '吸う', '触る', '愛撫',
            # Romanized terms
            'ecchi', 'oppai', 'sex', 'kiss', 'naked', 'breast',
            # English terms commonly found
            'arousal', 'climax', 'orgasm', 'penetration', 'intimate',
            'sensual', 'erotic', 'passionate', 'desire', 'lust',
            # Context indicators
            'moaning', 'breathing heavily', 'heartbeat', 'trembling',
            'hot breath', 'flushed', 'pleasure', 'ecstasy'
        ]
        
        # Initialize tokenizer for chunk splitting
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except:
            self.tokenizer = None
            logging.warning("tiktoken not available. Using character-based chunking.")
    
    def detect_eroge_content(self, text: str, threshold: float = 0.02) -> bool:
        """
        Detect if text contains erotic content based on keyword frequency.
        
        Args:
            text: Text to analyze
            threshold: Minimum ratio of eroge keywords to total words (default 2%)
            
        Returns:
            True if erotic content detected
        """
        if not text:
            return False
            
        text_lower = text.lower()
        total_words = len(text.split())
        
        if total_words == 0:
            return False
            
        eroge_count = 0
        for keyword in self.eroge_keywords:
            eroge_count += text_lower.count(keyword.lower())
        
        ratio = eroge_count / total_words
        return ratio >= threshold
    
    def split_text_into_chunks(self, text: str, max_tokens: int = 1500, 
                              overlap_tokens: int = 100) -> List[str]:
        """
        Split text into chunks with smart sentence boundary detection.
        
        Args:
            text: Text to split
            max_tokens: Maximum tokens per chunk
            overlap_tokens: Overlap between chunks for context
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
            
        # If no tokenizer available, use character-based approximation
        if not self.tokenizer:
            # Rough approximation: 1 token ≈ 4 characters for Japanese
            max_chars = max_tokens * 4
            overlap_chars = overlap_tokens * 4
            return self._split_by_characters(text, max_chars, overlap_chars)
        
        # Split into sentences first
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = len(self.tokenizer.encode(sentence))
            
            # If adding this sentence would exceed the limit
            if current_tokens + sentence_tokens > max_tokens and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, overlap_tokens)
                current_chunk = overlap_text + sentence
                current_tokens = len(self.tokenizer.encode(current_chunk))
            else:
                current_chunk += sentence
                current_tokens += sentence_tokens
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences, handling Japanese and English punctuation"""
        # Japanese sentence endings: 。！？
        # English sentence endings: . ! ?
        sentence_pattern = r'[。！？.!?]+\s*'
        sentences = re.split(sentence_pattern, text)
        
        # Rejoin sentence endings
        result = []
        matches = re.finditer(sentence_pattern, text)
        endings = [match.group() for match in matches]
        
        for i, sentence in enumerate(sentences):
            if sentence.strip():  # Skip empty sentences
                if i < len(endings):
                    result.append(sentence + endings[i])
                else:
                    result.append(sentence)
        
        return result
    
    def _get_overlap_text(self, text: str, overlap_tokens: int) -> str:
        """Get the last portion of text for overlap"""
        if not self.tokenizer:
            # Character-based approximation
            overlap_chars = overlap_tokens * 4
            return text[-overlap_chars:] if len(text) > overlap_chars else ""
        
        tokens = self.tokenizer.encode(text)
        if len(tokens) <= overlap_tokens:
            return text
        
        overlap_token_slice = tokens[-overlap_tokens:]
        return self.tokenizer.decode(overlap_token_slice)
    
    def _split_by_characters(self, text: str, max_chars: int, overlap_chars: int) -> List[str]:
        """Fallback character-based splitting when tokenizer unavailable"""
        if len(text) <= max_chars:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chars
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to find a good break point (sentence ending)
            search_start = max(start, end - 200)  # Look back up to 200 chars
            break_points = []
            
            for punct in ['。', '！', '？', '.', '!', '?']:
                pos = text.rfind(punct, search_start, end)
                if pos != -1:
                    break_points.append(pos + 1)
            
            if break_points:
                end = max(break_points)
            
            chunks.append(text[start:end])
            start = end - overlap_chars
        
        return chunks
    
    def create_specialized_prompt(self, text: str, is_eroge: bool = False) -> str:
        """Create specialized prompt for light novel translation"""
        base_prompt = """You are a professional translator specializing in Japanese light novels. 
Translate the following text to English while maintaining:
- Natural, flowing English prose
- Character personality and tone
- Cultural context and nuances
- Proper formatting and line breaks

"""
        
        if is_eroge:
            eroge_addition = """This text contains mature/erotic content. Translate it naturally while:
- Using appropriate mature language when needed
- Maintaining the sensual and emotional tone
- Preserving the intimate atmosphere
- Using tasteful but accurate terminology

"""
            return base_prompt + eroge_addition
        
        return base_prompt
    
    def create_specialized_vocabulary(self, is_eroge: bool = False) -> str:
        """Create specialized vocabulary for light novel translation"""
        base_vocab = """Common Light Novel Terms:
先輩 (senpai) - senior, upperclassman
後輩 (kouhai) - junior, underclassman  
部活 (bukatsu) - club activities
生徒会 (seitokai) - student council
文化祭 (bunkasai) - cultural festival
体育祭 (taiikusai) - sports festival
教室 (kyoushitsu) - classroom
廊下 (rouka) - hallway
屋上 (okujou) - rooftop
保健室 (hokenshitsu) - infirmary
図書館 (toshokan) - library

"""
        
        if is_eroge:
            eroge_vocab = """Mature Content Terms:
愛撫 (aibu) - caress, intimate touch
吐息 (toiki) - sigh, breath
鼓動 (kodou) - heartbeat
体温 (taion) - body temperature
肌 (hada) - skin
唇 (kuchibiru) - lips
瞳 (hitomi) - eyes/pupils
髪 (kami) - hair
香り (kaori) - scent, fragrance
温もり (nukumori) - warmth

"""
            return base_vocab + eroge_vocab
            
        return base_vocab
        
    def _check_docx_support(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            logging.warning("python-docx not available. DOCX support disabled.")
            return False
    
    def _check_pdf_support(self) -> bool:
        """Check if PyPDF2/pdfplumber is available"""
        try:
            import PyPDF2
            import pdfplumber
            return True
        except ImportError:
            logging.warning("PyPDF2/pdfplumber not available. PDF support disabled.")
            return False
    
    def _check_epub_support(self) -> bool:
        """Check if ebooklib is available"""
        try:
            import ebooklib
            from ebooklib import epub
            return True
        except ImportError:
            logging.warning("ebooklib not available. EPUB support disabled.")
            return False
    
    def can_process(self, file_path: str) -> bool:
        """Check if file can be processed"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return True
        elif ext == '.docx':
            return self.docx_available
        elif ext == '.pdf':
            return self.pdf_available
        elif ext == '.epub':
            return self.epub_available
        
        return False
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from light novel file"""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.txt':
                return self._extract_from_txt(file_path)
            elif ext == '.docx' and self.docx_available:
                return self._extract_from_docx(file_path)
            elif ext == '.pdf' and self.pdf_available:
                return self._extract_from_pdf(file_path)
            elif ext == '.epub' and self.epub_available:
                return self._extract_from_epub(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return {"error": str(e), "chapters": [], "metadata": {}}
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result['encoding'] or 'utf-8'
        
        # Read file content
        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            content = f.read()
        
        # Split into chapters (look for chapter markers)
        chapters = self._split_into_chapters(content)
        
        return {
            "chapters": chapters,
            "metadata": {
                "format": "txt",
                "encoding": encoding,
                "total_chapters": len(chapters),
                "file_size": os.path.getsize(file_path)
            }
        }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        import docx
        
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Combine paragraphs and split into chapters
        content = '\n'.join(paragraphs)
        chapters = self._split_into_chapters(content)
        
        # Extract metadata
        props = doc.core_properties
        metadata = {
            "format": "docx",
            "title": props.title or "",
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
            "total_chapters": len(chapters),
            "paragraphs": len(paragraphs)
        }
        
        return {
            "chapters": chapters,
            "metadata": metadata
        }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        import pdfplumber
        
        text_pages = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    text_pages.append({
                        "page": page_num + 1,
                        "text": text.strip()
                    })
        
        # Combine all pages
        content = '\n'.join([page["text"] for page in text_pages])
        chapters = self._split_into_chapters(content)
        
        metadata = {
            "format": "pdf",
            "total_pages": len(text_pages),
            "total_chapters": len(chapters)
        }
        
        return {
            "chapters": chapters,
            "metadata": metadata,
            "pages": text_pages
        }
    
    def _extract_from_epub(self, file_path: str) -> Dict[str, Any]:
        """Extract text from EPUB file with chapter support"""
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        
        book = epub.read_epub(file_path)
        
        chapters = []
        chapter_count = 0
        
        # Extract chapters from spine
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Parse HTML content
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text content
                text = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                if text and len(text) > 100:  # Skip very short content
                    chapter_count += 1
                    
                    # Try to extract chapter title
                    title_elem = soup.find(['h1', 'h2', 'h3', 'title'])
                    chapter_title = title_elem.get_text().strip() if title_elem else f"Chapter {chapter_count}"
                    
                    chapters.append({
                        "id": f"chapter_{chapter_count}",
                        "title": chapter_title,
                        "content": text,
                        "word_count": len(text.split()),
                        "file_name": item.get_name()
                    })
        
        # Extract metadata
        metadata = {
            "format": "epub",
            "title": book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else "",
            "author": book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else "",
            "language": book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else "",
            "publisher": book.get_metadata('DC', 'publisher')[0][0] if book.get_metadata('DC', 'publisher') else "",
            "total_chapters": len(chapters)
        }
        
        return {
            "chapters": chapters,
            "metadata": metadata
        }
    
    def _split_into_chapters(self, content: str) -> List[Dict[str, Any]]:
        """Split content into chapters based on common patterns"""
        
        # Common chapter patterns
        chapter_patterns = [
            r'第\s*[0-9一二三四五六七八九十百千]+\s*[章话回節]',  # Japanese: 第X章, 第X話, etc.
            r'Chapter\s+\d+',  # English: Chapter 1, Chapter 2, etc.
            r'^\s*\d+\s*$',  # Simple numbers on their own line
            r'[★☆◆◇■□▲△●○]\s*.*?[★☆◆◇■□▲△●○]',  # Symbols
            r'={3,}.*?={3,}',  # Equal signs
            r'-{3,}.*?-{3,}',  # Dashes
        ]
        
        # Try to find chapter breaks
        chapter_breaks = []
        for pattern in chapter_patterns:
            matches = list(re.finditer(pattern, content, re.MULTILINE | re.IGNORECASE))
            if matches:
                chapter_breaks = [(m.start(), m.end(), m.group()) for m in matches]
                break
        
        if not chapter_breaks:
            # If no chapter patterns found, split by length
            return self._split_by_length(content)
        
        chapters = []
        
        for i, (start, end, title) in enumerate(chapter_breaks):
            # Get content from current chapter break to next one
            if i + 1 < len(chapter_breaks):
                next_start = chapter_breaks[i + 1][0]
                chapter_content = content[end:next_start].strip()
            else:
                # Last chapter
                chapter_content = content[end:].strip()
            
            if chapter_content:
                chapters.append({
                    "id": f"chapter_{i + 1}",
                    "title": title.strip(),
                    "content": chapter_content,
                    "word_count": len(chapter_content.split())
                })
        
        return chapters
    
    def _split_by_length(self, content: str, max_length: int = 5000) -> List[Dict[str, Any]]:
        """Split content by length when no chapter markers are found"""
        chapters = []
        words = content.split()
        
        current_chunk = []
        chapter_num = 1
        
        for word in words:
            current_chunk.append(word)
            
            if len(' '.join(current_chunk)) >= max_length:
                chapters.append({
                    "id": f"section_{chapter_num}",
                    "title": f"Section {chapter_num}",
                    "content": ' '.join(current_chunk),
                    "word_count": len(current_chunk)
                })
                current_chunk = []
                chapter_num += 1
        
        # Add remaining content
        if current_chunk:
            chapters.append({
                "id": f"section_{chapter_num}",
                "title": f"Section {chapter_num}",
                "content": ' '.join(current_chunk),
                "word_count": len(current_chunk)
            })
        
        return chapters
    
    def get_translatable_content(self, extracted_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Get translatable content from extracted data"""
        translatable_content = []
        
        if "error" in extracted_data:
            return translatable_content
        
        chapters = extracted_data.get("chapters", [])
        
        for chapter in chapters:
            if isinstance(chapter, dict) and "content" in chapter:
                # Split content into paragraphs for better translation
                paragraphs = self._split_into_paragraphs(chapter["content"])
                
                for i, paragraph in enumerate(paragraphs):
                    if self._contains_japanese(paragraph):
                        translatable_content.append({
                            "id": f"{chapter['id']}_para_{i + 1}",
                            "chapter_id": chapter["id"],
                            "chapter_title": chapter.get("title", ""),
                            "original": paragraph,
                            "context": "light_novel_paragraph",
                            "paragraph_number": i + 1
                        })
        
        return translatable_content
    
    def _split_into_paragraphs(self, content: str) -> List[str]:
        """Split content into paragraphs"""
        # Split by double newlines or specific patterns
        paragraphs = re.split(r'\n\s*\n|\r\n\s*\r\n', content)
        
        # Clean up paragraphs
        cleaned = []
        for para in paragraphs:
            para = para.strip()
            if para and len(para) > 10:  # Skip very short paragraphs
                cleaned.append(para)
        
        return cleaned
    
    def _contains_japanese(self, text: str) -> bool:
        """Check if text contains Japanese characters"""
        japanese_pattern = re.compile(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]')
        return bool(japanese_pattern.search(text))
    
    def create_translation_file(self, output_path: str, translated_content: List[Dict[str, Any]], 
                              original_metadata: Dict[str, Any], output_format: str = 'txt') -> str:
        """Create translated file in specified format"""
        
        if output_format == 'txt':
            return self._create_txt_output(output_path, translated_content, original_metadata)
        elif output_format == 'docx' and self.docx_available:
            return self._create_docx_output(output_path, translated_content, original_metadata)
        elif output_format == 'pdf' and self.pdf_available:
            return self._create_pdf_output(output_path, translated_content, original_metadata)
        elif output_format == 'epub' and self.epub_available:
            return self._create_epub_output(output_path, translated_content, original_metadata)
        else:
            # Fallback to TXT
            return self._create_txt_output(output_path, translated_content, original_metadata)
    
    def _create_txt_output(self, output_path: str, translated_content: List[Dict[str, Any]], 
                          metadata: Dict[str, Any]) -> str:
        """Create TXT output file"""
        
        # Group content by chapters
        chapters = {}
        for item in translated_content:
            chapter_id = item.get("chapter_id", "unknown")
            if chapter_id not in chapters:
                chapters[chapter_id] = {
                    "title": item.get("chapter_title", ""),
                    "paragraphs": []
                }
            chapters[chapter_id]["paragraphs"].append(item.get("translated", item.get("original", "")))
        
        # Write to file
        output_file = output_path.replace('.txt', '_translated.txt')
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Write metadata as header
            f.write(f"=== Translated Light Novel ===\n")
            if metadata.get("title"):
                f.write(f"Title: {metadata['title']}\n")
            if metadata.get("author"):
                f.write(f"Author: {metadata['author']}\n")
            f.write(f"Total Chapters: {metadata.get('total_chapters', len(chapters))}\n")
            f.write(f"Original Format: {metadata.get('format', 'unknown')}\n")
            f.write("\n" + "="*50 + "\n\n")
            
            # Write chapters
            for chapter_id in sorted(chapters.keys()):
                chapter = chapters[chapter_id]
                f.write(f"{chapter['title']}\n")
                f.write("-" * len(chapter['title']) + "\n\n")
                
                for paragraph in chapter['paragraphs']:
                    f.write(f"{paragraph}\n\n")
                
                f.write("\n")
        
        return output_file
    
    def _create_docx_output(self, output_path: str, translated_content: List[Dict[str, Any]], 
                           metadata: Dict[str, Any]) -> str:
        """Create DOCX output file"""
        import docx
        from docx.shared import Inches
        
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading(metadata.get('title', 'Translated Light Novel'), 0)
        
        # Add metadata
        if metadata.get('author'):
            doc.add_paragraph(f"Author: {metadata['author']}")
        doc.add_paragraph(f"Total Chapters: {metadata.get('total_chapters', 'Unknown')}")
        doc.add_paragraph(f"Original Format: {metadata.get('format', 'unknown').upper()}")
        
        doc.add_page_break()
        
        # Group content by chapters
        chapters = {}
        for item in translated_content:
            chapter_id = item.get("chapter_id", "unknown")
            if chapter_id not in chapters:
                chapters[chapter_id] = {
                    "title": item.get("chapter_title", ""),
                    "paragraphs": []
                }
            chapters[chapter_id]["paragraphs"].append(item.get("translated", item.get("original", "")))
        
        # Add chapters
        for chapter_id in sorted(chapters.keys()):
            chapter = chapters[chapter_id]
            
            # Chapter title
            doc.add_heading(chapter['title'], level=1)
            
            # Chapter content
            for paragraph in chapter['paragraphs']:
                doc.add_paragraph(paragraph)
        
        # Save file
        output_file = output_path.replace('.docx', '_translated.docx')
        doc.save(output_file)
        
        return output_file
    
    def _create_pdf_output(self, output_path: str, translated_content: List[Dict[str, Any]], 
                          metadata: Dict[str, Any]) -> str:
        """Create PDF output file"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        output_file = output_path.replace('.pdf', '_translated.pdf')
        
        # Create document
        doc = SimpleDocTemplate(output_file, pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading_style = styles['Heading1']
        normal_style = styles['Normal']
        
        # Story list
        story = []
        
        # Add title
        story.append(Paragraph(metadata.get('title', 'Translated Light Novel'), title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata
        if metadata.get('author'):
            story.append(Paragraph(f"Author: {metadata['author']}", normal_style))
        story.append(Paragraph(f"Total Chapters: {metadata.get('total_chapters', 'Unknown')}", normal_style))
        story.append(Paragraph(f"Original Format: {metadata.get('format', 'unknown').upper()}", normal_style))
        story.append(PageBreak())
        
        # Group content by chapters
        chapters = {}
        for item in translated_content:
            chapter_id = item.get("chapter_id", "unknown")
            if chapter_id not in chapters:
                chapters[chapter_id] = {
                    "title": item.get("chapter_title", ""),
                    "paragraphs": []
                }
            chapters[chapter_id]["paragraphs"].append(item.get("translated", item.get("original", "")))
        
        # Add chapters
        for chapter_id in sorted(chapters.keys()):
            chapter = chapters[chapter_id]
            
            # Chapter title
            story.append(Paragraph(chapter['title'], heading_style))
            story.append(Spacer(1, 12))
            
            # Chapter content
            for paragraph in chapter['paragraphs']:
                story.append(Paragraph(paragraph, normal_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        return output_file
    
    def _create_epub_output(self, output_path: str, translated_content: List[Dict[str, Any]], 
                           metadata: Dict[str, Any]) -> str:
        """Create EPUB output file"""
        import ebooklib
        from ebooklib import epub
        
        output_file = output_path.replace('.epub', '_translated.epub')
        
        # Create book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier('translated_novel')
        book.set_title(metadata.get('title', 'Translated Light Novel'))
        book.set_language('en')  # Target language
        
        if metadata.get('author'):
            book.add_author(metadata['author'])
        
        # Group content by chapters
        chapters = {}
        for item in translated_content:
            chapter_id = item.get("chapter_id", "unknown")
            if chapter_id not in chapters:
                chapters[chapter_id] = {
                    "title": item.get("chapter_title", ""),
                    "paragraphs": []
                }
            chapters[chapter_id]["paragraphs"].append(item.get("translated", item.get("original", "")))
        
        # Create chapters
        epub_chapters = []
        for i, chapter_id in enumerate(sorted(chapters.keys())):
            chapter = chapters[chapter_id]
            
            # Create chapter content
            content = f"<h1>{chapter['title']}</h1>\n"
            for paragraph in chapter['paragraphs']:
                content += f"<p>{paragraph}</p>\n"
            
            # Create EPUB chapter
            epub_chapter = epub.EpubHtml(title=chapter['title'],
                                       file_name=f'chapter_{i+1}.xhtml',
                                       lang='en')
            epub_chapter.content = content
            
            book.add_item(epub_chapter)
            epub_chapters.append(epub_chapter)
        
        # Define Table of Contents
        book.toc = epub_chapters
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Cambria, Liberation Serif, Bitstream Vera Serif, Georgia, Times, Times New Roman, serif;
        }
        h1 {
            text-align: left;
            text-transform: uppercase;
            font-weight: 200;
        }
        '''
        
        nav_css = epub.EpubItem(uid="nav_css", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Create spine
        book.spine = ['nav'] + epub_chapters
        
        # Write EPUB file
        epub.write_epub(output_file, book, {})
        
        return output_file

    def get_light_novel_prompt(self) -> str:
        """Get specialized prompt for light novel translation"""
        return """You are a professional light novel translator specializing in Japanese to English translation. Your task is to translate Japanese light novel text while maintaining the narrative flow, character voice, and literary quality.

Guidelines for light novel translation:
1. **Narrative Style**: Preserve the author's narrative voice and tone
2. **Character Voice**: Maintain distinct character personalities and speech patterns
3. **Literary Flow**: Ensure smooth, natural English that reads well as a novel
4. **Cultural Context**: Adapt cultural references appropriately for English readers
5. **Consistency**: Keep character names, terms, and world-building elements consistent
6. **Readability**: Prioritize readability and engagement over literal accuracy

Translation approach:
- Use natural, flowing English prose suitable for novel reading
- Adapt Japanese speech patterns to English equivalents that convey the same tone
- Maintain paragraph structure and pacing
- Keep honorifics only when they add meaningful context
- Translate sound effects and onomatopoeia creatively
- Preserve the emotional impact and atmosphere of the original

Remember: You are creating an English light novel that should read naturally while staying true to the original story and characters."""

    def get_light_novel_vocabulary(self) -> List[str]:
        """Get specialized vocabulary for light novel translation"""
        return [
            "先輩 (senpai) - senior/upperclassman",
            "後輩 (kouhai) - junior/underclassman", 
            "先生 (sensei) - teacher/master",
            "お兄ちゃん (onii-chan) - big brother (affectionate)",
            "お姉ちゃん (onee-chan) - big sister (affectionate)",
            "お父さん (otou-san) - father",
            "お母さん (okaa-san) - mother",
            "学園 (gakuen) - academy/school",
            "部活 (bukkatsu) - club activities",
            "文化祭 (bunka-sai) - cultural festival",
            "体育祭 (taiiku-sai) - sports festival",
            "修学旅行 (shuugaku-ryokou) - school trip",
            "クラスメート (classmate) - classmate",
            "友達 (tomodachi) - friend",
            "恋人 (koibito) - lover/boyfriend/girlfriend",
            "告白 (kokuhaku) - confession (of love)",
            "デート (date) - date",
            "バレンタイン (Valentine) - Valentine's Day",
            "ホワイトデー (White Day) - White Day",
            "クリスマス (Christmas) - Christmas",
            "お正月 (oshougatsu) - New Year",
            "夏休み (natsuyasumi) - summer vacation",
            "冬休み (fuyuyasumi) - winter vacation",
            "春休み (haruyasumi) - spring vacation",
            "桜 (sakura) - cherry blossoms",
            "花見 (hanami) - cherry blossom viewing",
            "制服 (seifuku) - school uniform",
            "お弁当 (obentou) - lunch box",
            "カフェテリア (cafeteria) - cafeteria",
            "図書館 (toshokan) - library",
            "屋上 (okujou) - rooftop",
            "教室 (kyoushitsu) - classroom",
            "廊下 (rouka) - hallway",
            "階段 (kaidan) - stairs",
            "昇降口 (shoukou-guchi) - school entrance",
            "下駄箱 (getabako) - shoe locker",
            "ロッカー (locker) - locker",
            "職員室 (shokuin-shitsu) - faculty room",
            "保健室 (hoken-shitsu) - nurse's office",
            "委員会 (iinkai) - committee",
            "生徒会 (seitokai) - student council",
            "会長 (kaichou) - president",
            "副会長 (fuku-kaichou) - vice president",
            "書記 (shoki) - secretary",
            "会計 (kaikei) - treasurer",
            "監査 (kansa) - auditor"
        ]
